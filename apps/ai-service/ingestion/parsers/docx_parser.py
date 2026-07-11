"""
DOCX parser using python-docx.
Groups paragraphs into pseudo-pages since .docx has no native page concept
(page breaks depend on rendering, not stored layout).
We batch every N paragraphs into one PageContent unit — keeps chunking
input consistent with PDF's page-based structure.
"""

from docx import Document
from ..schema import IngestResult, PageContent, FileType

PARAGRAPHS_PER_BLOCK = 15  # tune based on avg paragraph length


def parse_docx(file_path: str, filename: str) -> IngestResult:
    doc = Document(file_path)

    # Include table content too — tables often carry key data (specs, pricing)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    pages = []
    for i in range(0, len(paragraphs), PARAGRAPHS_PER_BLOCK):
        block = paragraphs[i : i + PARAGRAPHS_PER_BLOCK]
        block_text = "\n".join(block)
        page_num = (i // PARAGRAPHS_PER_BLOCK) + 1
        pages.append(PageContent(page_num=page_num, text=block_text))

    is_empty = len(pages) == 0 or sum(p.char_count for p in pages) < 20

    return IngestResult(
        filename=filename,
        file_type=FileType.DOCX,
        pages=pages,
        is_scanned_or_empty=is_empty,
        warning="Document appears empty or unreadable." if is_empty else None,
    )
